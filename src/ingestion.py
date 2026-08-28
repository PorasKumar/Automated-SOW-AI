import email
from email import policy
import docx
from pypdf import PdfReader
from pathlib import Path
from typing import List,Tuple
import io

def extract_text_from_file(file_bytes:bytes, filename:str) -> str:
    """Will extract texts from different types of files, so that we can structure it in json format in next phase"""

    #if no data in file, then skip
    if not file_bytes:
        print(f"[PARSER WARNING] File '{filename}' contained 0 bytes. Skipping.")
        return ""
    
    try:
        ext = filename.lower().split(".")[-1]

        # Create stream buffer from raw bytes, so that we can read docx and pypdf without any issue
        file_stream = io.BytesIO(file_bytes)

        if (ext=="eml"):
            msg = email.message_from_bytes(file_bytes, policy=policy.default)
            body = msg.get_body(preferencelist=("plain","html"))
            return body.get_content() if body else ""

        elif (ext=="pdf"):
            file_stream.seek(0)  # Reset stream cursor to beginning
            try:
                # strict=False prevents crashes on minor stream EndOfFile(EOF) errors
                reader = PdfReader(file_stream, strict=False)
                pages_text = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)

                return "\n".join(pages_text)

            except Exception as pdf_err:
                print(
                    f"[PDF PARSER ERROR] Failed reading PDF '{filename}': {pdf_err}"
                )
                return ""

        elif (ext == "docx"):
            file_stream.seek(0) # Reset stream cursor to beginning, kyuki agar ek file ko read karke end pe pohoch gaya, then next file keliye beginning pe laao
            doc = docx.Document(file_stream)
            text_blocks = []

            # Extract paragraph text
            for p in doc.paragraphs:
                if p.text.strip():
                    text_blocks.append(p.text.strip())

            # Extract table text (Crucial for SOW scope/pricing tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if row_text:
                        text_blocks.append(" | ".join(row_text))

            return "\n".join(text_blocks)
        

        elif ext in ["json", "txt", "md"]:
            return file_bytes.decode("utf-8")

        else:
            raise ValueError(f"Unsupported File Format .{ext}")

    except Exception as e:
            print(f"Error in file ingestion {e}")
            raise RuntimeError(f"Error in file ingestion {e}")


def uploaded_file_aggregator(files: List[Tuple[bytes, str]]) -> str:
    """Combines multiple extracted file streams into a delimited context string."""

    combined_context = []

    #no try catch block, because even if 1 of many file causes error, entire pipeline fails without reading other files

    for idx, (file_bytes, filename) in enumerate(files, start=1):
        print(f"Parsing file {idx}/{len(files)}: {filename}")
        extracted_text = extract_text_from_file(file_bytes, filename)

        if extracted_text.strip():
            file_block = (
                f"---Start Source File {idx}: {filename}---\n"
                f"{extracted_text.strip()}\n"
                f"---End Source File {idx}: {filename}---"
            )
            combined_context.append(file_block)

    return "\n\n".join(combined_context)